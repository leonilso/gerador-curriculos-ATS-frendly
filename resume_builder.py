from docx import Document
from llm_adapter import adapt_summary, adapt_objective

def build_docx(cv, match):
    doc = Document()

    summary = adapt_summary(cv, match)
    objective = adapt_objective(cv, match)

    core = doc.core_properties
    core.title = f"{match['title']}"
    core.author = f"{cv["profile"]["name"]}"
    core.keywords = (", ".join(match["matched_technologies"]))
    core.subject = objective
    core.language = "pt-BR"
    core.category = "Curriculo"
    core.comments = (
        "Auto-generated CV | "
        f"Target role: {match['title']} | "
        "ATS-friendly layout | "
        "v1.0"
    )

    # ===== CABEÇALHO =====
    p = doc.add_paragraph()
    p.add_run(cv["profile"]["name"]).bold = True
    p.add_run(f"\n{match['title']}")
    p.add_run(f"\ne-mail: {cv["profile"]['email']}")
    p.add_run(f"\ntelefone: {cv["profile"]['phone']}")
    for key, url in cv["profile"]['urls'].items():
        p.add_run(f"\n{key}: {url}")

    # ===== OBJETIVO =====
    doc.add_heading("Objetivo Profissional", level=2)
    doc.add_paragraph(objective)

    # ===== RESUMO =====
    doc.add_heading("Resumo Profissional", level=2)
    doc.add_paragraph(summary)

    # ===== COMPETÊNCIAS =====
    doc.add_heading("Competências Técnicas", level=2) 
    for skill in match["matched_hard_skills"]: 
        doc.add_paragraph(f"• {skill}")

    # ===== COMPETÊNCIAS: tecnologias =====
    if match["matched_technologies"]:
        p = doc.add_paragraph()
        p.add_run("Tecnologias: ").bold = True
        p.add_run(", ".join(match["matched_technologies"]))

    # ===== COMPETÊNCIAS: ferramentas =====
    if match["matched_tools"]:
        p = doc.add_paragraph()
        p.add_run("Ferramentas: " ).bold = True
        p.add_run(", ".join(match["matched_tools"]))

    # ===== FORMAÇÃO =====
    doc.add_heading("Formação Acadêmica", level=2)

    for edu in cv["education"]:
        period = (
            f"{edu['startDate']} – {edu['endDate']}"
            if "startDate" in edu and "endDate" in edu
            else edu.get("date", "")
        )

        doc.add_paragraph(
            f"{edu['area']} — {edu['institution']} ({period})"
        )

        if edu.get("description"):
            doc.add_paragraph(edu["description"])

    # ===== EXPERIÊNCIA =====
    doc.add_heading("Experiência Profissional", level=2)

    for exp in cv["experience"]:
        p = doc.add_paragraph()
        p.add_run(f"• {exp['position']}").bold = True
        p.add_run(f"— {exp['company']}\n" f"{exp['description']}")


    # ===== CERTIFICAÇÕES =====
    if match["matched_certifications"]:
        doc.add_heading("Certificações Relevantes", level=2)

        certs_line = "; ".join(
            f"{c['name']} ({c['issuer']}, {c['date']})"
            for c in match["matched_certifications"]
        )
        doc.add_paragraph(certs_line)

    # ===== PROJETOS =====
    if match["matched_projects"]:
        doc.add_heading("Projetos Relevantes", level=2)

        for p in match["matched_projects"]:
            par = doc.add_paragraph()
            par.add_run(f"{p['title']}").bold = True
            par.add_run(f" - Tecnologias: {', '.join(p['techStack'])} - {p['description']}")

            if p.get("url"):
                doc.add_paragraph(p["url"])

    # ===== IDIOMAS =====
    doc.add_heading("Idiomas", level=2)

    langs_line = "; ".join(
        f"{l['language']} ({l['fluency']})"
        for l in cv["skills"]["languages"]
    )
    doc.add_paragraph(langs_line)

    output = f"curriculos/curriculo_{match['title'].replace("/", "").replace(".", "")}.docx"
    doc.save(output)
